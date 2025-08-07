from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import markdown
import markdown2
from playwright.async_api import async_playwright
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import HtmlFormatter
import asyncio
from docx import Document
from docx.shared import Inches
import tempfile
import os
from typing import Optional
import uuid
import re
from markdown.extensions import codehilite, toc, tables
import bleach

app = FastAPI(title="Markdown to PDF Converter", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class StyleSettings(BaseModel):
    font_family: Optional[str] = None  # CSS font-family string
    h1_size: Optional[int] = None      # px
    h2_size: Optional[int] = None      # px
    h3_size: Optional[int] = None      # px
    p_size: Optional[int] = None       # px


class MarkdownRequest(BaseModel):
    content: str
    title: Optional[str] = "Document"
    settings: Optional[StyleSettings] = None

class ConversionResponse(BaseModel):
    success: bool
    message: str
    download_url: Optional[str] = None

# Store generated files temporarily
temp_files = {}

def normalize_markdown_bullets(content: str) -> str:
    """Normalize list markers and insert required blank lines after headings.

    - Converts leading en/em dashes and bullet-like glyphs to a standard "- "
    - Ensures there is a space after a leading dash
    - Inserts a blank line between a heading and an immediately following list
    - Skips transformations inside fenced code blocks
    """
    if not content:
        return content

    lines = content.splitlines()
    normalized_lines: list[str] = []
    inside_code_fence = False
    previous_raw_line = ""

    fence_pattern = re.compile(r"^\s*(```|~~~)")
    heading_pattern = re.compile(r"^\s*#{1,6}\s+")
    list_start_pattern = re.compile(r"^\s*([-*+]\s|\d+\.\s)")

    for raw_line in lines:
        line = raw_line

        # Toggle code fence blocks
        if fence_pattern.match(raw_line):
            inside_code_fence = not inside_code_fence
            normalized_lines.append(line)
            previous_raw_line = raw_line
            continue

        if not inside_code_fence:
            # Convert unconventional bullets to '- '
            if re.match(r"^\s*[–—•∙·‣]", line):
                line = re.sub(r"^(\s*)[–—•∙·‣]\s*", r"\1- ", line, count=1)

            # Ensure a space after a leading dash
            line = re.sub(r"^(\s*)-(?!\s)(.*)$", r"\1- \2", line)

            # If a heading is immediately followed by a list, insert a blank line
            if previous_raw_line and heading_pattern.match(previous_raw_line) and list_start_pattern.match(line):
                if normalized_lines and normalized_lines[-1] != "":
                    normalized_lines.append("")

        normalized_lines.append(line)
        previous_raw_line = raw_line

    return "\n".join(normalized_lines)

@app.get("/")
async def root():
    return {"message": "Markdown to PDF Converter API"}

@app.post("/api/convert/html", response_model=dict)
async def convert_to_html(request: MarkdownRequest):
    """Convert markdown to HTML for preview"""
    try:
        # Configure markdown with extensions
        md = markdown.Markdown(
            extensions=[
                'codehilite',
                'toc',
                'tables',
                'fenced_code',
                'nl2br'
            ],
            extension_configs={
                'codehilite': {
                    'css_class': 'highlight',
                    'use_pygments': True
                }
            }
        )
        
        # Normalize bullets before conversion
        normalized = normalize_markdown_bullets(request.content)
        # Convert markdown to HTML
        html_content = md.convert(normalized)
        
        # Clean HTML for security
        allowed_tags = [
            'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
            'p', 'br', 'strong', 'em', 'u', 's',
            'ul', 'ol', 'li', 'blockquote',
            'code', 'pre', 'div', 'span',
            'table', 'thead', 'tbody', 'tr', 'th', 'td',
            'a', 'img'
        ]
        
        allowed_attributes = {
            'a': ['href', 'title'],
            'img': ['src', 'alt', 'title'],
            'div': ['class'],
            'span': ['class'],
            'pre': ['class'],
            'code': ['class']
        }
        
        clean_html = bleach.clean(
            html_content,
            tags=allowed_tags,
            attributes=allowed_attributes,
            strip=True
        )
        
        return {
            "success": True,
            "html": clean_html,
            "title": request.title
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"HTML conversion failed: {str(e)}")

@app.post("/api/convert/pdf")
async def convert_to_pdf(request: MarkdownRequest):
    """Convert markdown to PDF in monochrome, mirroring web preview layout/fonts."""
    try:
        # Convert markdown to HTML with all extensions
        md = markdown.Markdown(
            extensions=[
                'codehilite',
                'toc',
                'tables',
                'fenced_code',
                'nl2br',
                'extra'
            ],
            extension_configs={
                'codehilite': {
                    'css_class': 'highlight',
                    'use_pygments': True,
                    'noclasses': True,
                    # Force monochrome style for PDF output
                    'pygments_style': 'bw'
                }
            }
        )
        
        # Normalize bullets first for consistent list detection
        normalized_md = normalize_markdown_bullets(request.content)
        html_content = md.convert(normalized_md)
        
        # Monochrome CSS mirroring frontend preview (fonts/layout), grayscale only
        css_styles = """
        <style>
            @page {
                size: A4;
                margin: 2cm;
                @bottom-right {
                    content: counter(page);
                    font-size: 10px;
                    color: #666;
                }
            }
            
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', sans-serif;
                line-height: 1.7;
                color: #111;
                font-size: 12px;
            }
            
            h1, h2, h3, h4, h5, h6 {
                color: #111;
                margin-top: 24px;
                margin-bottom: 12px;
                page-break-after: avoid;
                font-weight: 600;
            }
            
            h1 {
                font-size: 24px;
                border-bottom: 2px solid #e5e7eb;
                padding-bottom: 10px;
                margin-bottom: 20px;
            }
            
            h2 {
                font-size: 20px;
                border-bottom: 1px solid #e5e7eb;
                padding-bottom: 8px;
                margin-bottom: 16px;
            }
            
            h3 {
                font-size: 16px;
                color: #111;
                margin-bottom: 12px;
            }
            
            h4 {
                font-size: 14px;
                color: #111;
                margin-bottom: 10px;
            }
            
            p {
                margin-bottom: 16px;
                text-align: left;
                color: #222;
            }
            
            /* Monochrome code blocks similar to preview */
            pre {
                background-color: #f5f7fa;
                color: #111;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding: 16px;
                margin: 16px 0;
                font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
                font-size: 12px;
                line-height: 1.6;
                overflow: hidden;
                page-break-inside: avoid;
            }
            
            
            code {
                background-color: #f3f4f6;
                padding: 2px 6px;
                border-radius: 4px;
                font-family: 'Monaco', 'Menlo', 'Ubuntu Mono', monospace;
                font-size: 12px;
                color: #111;
            }
            
            /* Monochrome tables like preview */
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 16px 0;
                font-size: 12px;
                page-break-inside: avoid;
                border: 1px solid #d1d5db;
            }
            
            th, td {
                border: 1px solid #d1d5db;
                padding: 12px;
                text-align: left;
                vertical-align: top;
            }
            
            th {
                background-color: #f9fafb;
                font-weight: 600;
                color: #111;
            }
            
            tr:nth-child(even) { background-color: #fafafa; }
            
            /* Lists */
            ul, ol {
                margin: 12px 0;
                padding-left: 24px;
            }
            
            li {
                margin-bottom: 6px;
                line-height: 1.5;
            }
            
            ul li {
                list-style-type: disc;
            }
            
            ol li {
                list-style-type: decimal;
            }
            
            /* Blockquotes */
            blockquote {
                border-left: 4px solid #d1d5db;
                margin: 16px 0;
                padding-left: 16px;
                color: #6b7280;
                font-style: italic;
                background-color: #f8f9fa;
                padding: 12px 16px;
                border-radius: 0 4px 4px 0;
            }
            
            /* Links */
            a {
                color: #111;
                text-decoration: none;
            }
            
            a:hover {
                text-decoration: underline;
            }
            
            /* Horizontal rules */
            hr {
                border: none;
                height: 2px;
                background: #e5e7eb;
                margin: 24px 0;
            }
            
            /* Syntax highlighting container (monochrome) */
            .highlight {
                background-color: #f3f4f6;
                border-radius: 6px;
                padding: 16px;
                margin: 16px 0;
                font-size: 12px;
                line-height: 1.6;
            }
            
            /* Page breaks */
            .page-break {
                page-break-before: always;
            }
        </style>
        """
        
        # Optional style overrides from request settings
        ff = (request.settings.font_family if request.settings and request.settings.font_family else "-apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', sans-serif")
        h1s = (request.settings.h1_size if request.settings and request.settings.h1_size else 24)
        h2s = (request.settings.h2_size if request.settings and request.settings.h2_size else 20)
        h3s = (request.settings.h3_size if request.settings and request.settings.h3_size else 16)
        ps = (request.settings.p_size if request.settings and request.settings.p_size else 12)

        # Note: escape CSS braces in f-string using double braces
        css_override = f"""
        <style>
            body{{ font-family: {ff}; font-size: {ps}px; }}
            h1{{ font-size: {h1s}px; }}
            h2{{ font-size: {h2s}px; }}
            h3{{ font-size: {h3s}px; }}
            pre, code, .highlight{{ font-size: {ps}px; }}
        </style>
        """

        # Create the complete HTML document
        html_doc = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{request.title}</title>
            {css_styles}
            {css_override}
        </head>
        <body>
            <div style="max-width: 900px; margin: 0 auto;">
                <h1>{request.title}</h1>
                {html_content}
            </div>
        </body>
        </html>
        """
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        temp_dir = tempfile.gettempdir()
        pdf_path = os.path.join(temp_dir, f"{file_id}.pdf")
        
        # Convert HTML to PDF using Playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            
            # Set the HTML content
            await page.set_content(html_doc)
            
            # Generate PDF with beautiful options
            await page.pdf(
                path=pdf_path,
                format='A4',
                margin={
                    'top': '2cm',
                    'right': '2cm',
                    'bottom': '2cm',
                    'left': '2cm'
                },
                print_background=True,
                prefer_css_page_size=True
            )
            
            await browser.close()
        
        # Store file info for download
        temp_files[file_id] = {
            'path': pdf_path,
            'filename': f"{request.title.replace(' ', '_')}.pdf"
        }
        
        return {
            "success": True,
            "message": "PDF generated successfully",
            "download_url": f"/api/download/{file_id}"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF conversion failed: {str(e)}")

@app.post("/api/convert/docx")
async def convert_to_docx(request: MarkdownRequest):
    """Convert markdown to DOCX"""
    try:
        # Convert markdown to HTML first
        md = markdown.Markdown(
            extensions=['tables', 'fenced_code', 'nl2br']
        )
        html_content = md.convert(request.content)
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(request.title, 0)
        
        # Simple markdown to docx conversion
        # This is a basic implementation - for production, consider using python-docx-template
        lines = request.content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # Code blocks
            elif line.startswith('```'):
                continue  # Skip code block markers for now
            # Regular paragraphs
            else:
                # Basic formatting
                paragraph = doc.add_paragraph()
                
                # Handle bold and italic (basic implementation)
                text = line
                if '**' in text:
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            paragraph.add_run(part)
                        else:
                            paragraph.add_run(part).bold = True
                else:
                    paragraph.add_run(text)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        temp_dir = tempfile.gettempdir()
        docx_path = os.path.join(temp_dir, f"{file_id}.docx")
        
        # Save document
        doc.save(docx_path)
        
        # Store file info for download
        temp_files[file_id] = {
            'path': docx_path,
            'filename': f"{request.title.replace(' ', '_')}.docx"
        }
        
        return {
            "success": True,
            "message": "DOCX generated successfully",
            "download_url": f"/api/download/{file_id}"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX conversion failed: {str(e)}")

@app.get("/api/download/{file_id}")
async def download_file(file_id: str):
    """Download generated file"""
    if file_id not in temp_files:
        raise HTTPException(status_code=404, detail="File not found")
    
    file_info = temp_files[file_id]
    file_path = file_info['path']
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")
    
    return FileResponse(
        path=file_path,
        filename=file_info['filename'],
        media_type='application/octet-stream'
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
